"""Extraction strategies for non-PDF file formats.

Covers DOCX, plain text, and a non-textual fallback. The
``SpreadsheetExtractor`` lives in ``spreadsheet.py``.

Each strategy emits an ordered element stream
(``ExtractionResult.elements``) and a single concatenated
``PageResult`` so downstream consumers that still read ``pages`` /
``full_text`` keep working without change.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING, cast

from womblex.ingest.elements import Cell, Element, ElementKind
from womblex.ingest.extract import (
    ExtractionMetadata,
    ExtractionResult,
    PageResult,
)

if TYPE_CHECKING:
    import fitz

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class DocxExtractor:
    """Extract elements from a Word document in document order.

    Paragraphs and tables are interleaved as they appear in the source
    body — required for budget-statement / portfolio-document style
    files where prose and tables alternate and downstream consumers
    care which table follows which paragraph.
    """

    def extract_path(self, path: Path) -> ExtractionResult:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            return _docx_error("python-docx not installed; cannot extract DOCX.")

        try:
            doc = Document(str(path))
        except Exception as e:
            return _docx_error(f"Failed to read DOCX: {e}")

        # python-docx returns paragraphs and tables in separate collections.
        # Look them up by their underlying XML element so body iteration can
        # walk paragraphs and tables in true document order.
        paras_by_xml = {p._element: p for p in doc.paragraphs}
        tables_by_xml = {t._element: t for t in doc.tables}

        elements: list[Element] = []
        order = 0
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                para = paras_by_xml.get(child)
                if para is None or not para.text.strip():
                    continue
                kind = cast("ElementKind", (
                    "heading"
                    if para.style and "heading" in para.style.name.lower()
                    else "paragraph"
                ))
                elements.append(Element(
                    order=order, kind=kind, extractor="docx",
                    text=para.text, confidence=0.9,
                ))
                order += 1
            elif child.tag == qn("w:tbl"):
                tbl = tables_by_xml.get(child)
                if tbl is None:
                    continue
                cells: list[Cell] = []
                for row_idx, row in enumerate(tbl.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cells.append(Cell(row=row_idx, col=col_idx, value=cell.text))
                header_rows = [0] if cells else []
                elements.append(Element(
                    order=order, kind="table", extractor="docx",
                    cells=cells, header_rows=header_rows, confidence=0.85,
                ))
                order += 1

        page_text = "\n\n".join(e.text for e in elements if e.text)
        return ExtractionResult(
            pages=[PageResult(page_number=0, text=page_text, method="docx")],
            elements=elements,
            method="docx",
            metadata=ExtractionMetadata(
                extraction_strategy="docx",
                confidence=0.9,
                processing_time=0.0,
                page_count=1,
                text_coverage=1.0 if page_text else 0.0,
            ),
        )


def _docx_error(msg: str) -> ExtractionResult:
    return ExtractionResult(
        pages=[], method="docx", error=msg,
        metadata=ExtractionMetadata(
            extraction_strategy="docx",
            confidence=0.0, processing_time=0.0,
            page_count=0, text_coverage=0.0,
        ),
    )


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


class TextExtractor:
    """Plain text passthrough. One paragraph element per blank-line-separated block."""

    def extract_path(self, path: Path) -> ExtractionResult:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")

        blocks = [b for b in text.split("\n\n") if b.strip()]
        elements = [
            Element(order=i, kind="paragraph", extractor="text", text=b, confidence=1.0)
            for i, b in enumerate(blocks)
        ]
        return ExtractionResult(
            pages=[PageResult(page_number=0, text=text, method="text")],
            elements=elements,
            method="text",
            metadata=ExtractionMetadata(
                extraction_strategy="text",
                confidence=1.0,
                processing_time=0.0,
                page_count=1,
                text_coverage=1.0 if text.strip() else 0.0,
            ),
        )


# ---------------------------------------------------------------------------
# Non-textual fallback
# ---------------------------------------------------------------------------


class NonTextualExtractor:
    """Placeholder for documents that cannot be extracted — flags for manual review."""

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        return ExtractionResult(
            pages=[], method="non_textual",
            error="Document flagged as non-textual; requires manual review.",
            metadata=ExtractionMetadata(
                extraction_strategy="non_textual",
                confidence=0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=0.0,
            ),
        )
