"""Extraction strategy implementations for each document type."""

from __future__ import annotations

import logging
from pathlib import Path

import fitz
import numpy as np

from womblex.ingest.extract import (
    ExtractionMetadata, ExtractionResult, ImageData, PageResult,
    Position, TableData, TextBlock,
    _build_text_blocks, _extract_form_fields, _extract_images_from_page,
    _extract_tables_from_page, _normalise_bbox, _ocr_text_block, _page_to_gray,
)
from womblex.ingest.paddle_ocr import (
    get_layout_analyzer, get_paddle_reader, get_table_recognizer,
    preprocess_for_ocr,
)

logger = logging.getLogger(__name__)


def _ocr_page(
    page: fitz.Page, dpi: int, lang: str,
) -> tuple[str, float, list[str]]:
    """OCR a page: blur check → deskew → binarise → PaddleOCR. Confidence 0-100."""
    import cv2

    pix = page.get_pixmap(dpi=dpi)
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

    # Pre-OCR blur check — flag pages too blurry for reliable OCR
    pre_gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY) if pix.n >= 3 else img
    from womblex.ingest.heuristics_cv2 import calculate_blur_score
    blur = calculate_blur_score(pre_gray)

    gray, steps = preprocess_for_ocr(img)

    if blur is not None and blur < 50:
        steps.append("low_blur_warning")
        logger.warning("blurry page: doc=%s page=%d blur_score=%.1f", page.parent.name, page.number, blur)

    reader = get_paddle_reader(lang)
    results = reader.readtext(gray)
    text = "\n".join(r[1] for r in results if r[1].strip())
    region_confs = [float(r[2]) for r in results if r[1].strip()]
    avg_conf = (sum(region_confs) / len(region_confs)) * 100 if region_confs else 0.0

    if avg_conf < 40.0:
        logger.warning("low OCR confidence: doc=%s page=%d confidence=%.1f", page.parent.name, page.number, avg_conf)

    return text, avg_conf, steps


def _layout_blocks_and_tables(
    page: fitz.Page,
    dpi: int,
    text: str,
    conf: float,
) -> tuple[list[TextBlock], list[TableData]]:
    """Run YOLO layout analysis on a page, returning typed TextBlocks and tables.

    Falls back to a single paragraph block if the layout model is unavailable.
    """
    blocks: list[TextBlock] = []
    tables: list[TableData] = []

    try:
        analyzer = get_layout_analyzer()
        pix = page.get_pixmap(dpi=dpi)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

        regions = analyzer.analyze(img)
        if not regions:
            raise RuntimeError("no layout regions detected")

        for region in regions:
            rx0, ry0, rx1, ry1 = region.bbox
            pos = _normalise_bbox((rx0, ry0, rx1, ry1), float(pix.width), float(pix.height))

            if region.block_type == "table":
                # Try SLANet table recognition on this region
                try:
                    recognizer = get_table_recognizer()
                    table_result = recognizer.recognize(img, region.bbox)
                    if table_result and (table_result.headers or table_result.rows):
                        tables.append(TableData(
                            headers=table_result.headers,
                            rows=table_result.rows,
                            position=pos,
                            confidence=region.confidence,
                        ))
                        continue
                except (FileNotFoundError, Exception):
                    pass
                # Fall through to text block if table recognition unavailable
                blocks.append(TextBlock(
                    text="[TABLE]",
                    position=pos,
                    block_type="table",
                    confidence=region.confidence,
                ))
            else:
                blocks.append(TextBlock(
                    text="",  # layout region text not yet segmented from OCR output
                    position=pos,
                    block_type=region.block_type,
                    confidence=region.confidence,
                ))

        # If layout produced blocks but none have text, fall back to single block
        if blocks and not any(b.text.strip() for b in blocks if b.block_type != "table"):
            block = _ocr_text_block(page, text, conf)
            if block:
                # Assign the layout-derived block_type from the dominant region
                dominant = max(regions, key=lambda r: (r.bbox[2] - r.bbox[0]) * (r.bbox[3] - r.bbox[1]))
                block = TextBlock(
                    text=block.text,
                    position=block.position,
                    block_type=dominant.block_type,
                    confidence=block.confidence,
                )
                return [block], tables

    except (FileNotFoundError, Exception):
        # Layout model not available — fall back to heuristic block
        pass

    if not blocks:
        block = _ocr_text_block(page, text, conf)
        if block:
            blocks = [block]

    return blocks, tables


def _text_coverage(pages: list[PageResult]) -> float:
    """Fraction of pages with meaningful text (>50 chars)."""
    if not pages:
        return 0.0
    filled = sum(1 for p in pages if len(p.text.strip()) > 50)
    return filled / len(pages)


# ---------------------------------------------------------------------------
# 1. native_narrative
# ---------------------------------------------------------------------------


class NativeNarrativeExtractor:
    """Extract clean text preserving paragraph structure from native PDFs."""

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_blocks: list[TextBlock] = []

        for page in doc:
            text = page.get_text("text", flags=fitz.TEXT_DEHYPHENATE)
            pages.append(PageResult(page_number=page.number, text=text, method="native"))
            all_blocks.extend(_build_text_blocks(page))

        coverage = _text_coverage(pages)
        return ExtractionResult(
            pages=pages,
            method="native_narrative",
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="native_narrative",
                confidence=min(0.95, 0.8 + coverage * 0.15),
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
            ),
        )


# ---------------------------------------------------------------------------
# 2. native_with_structured
# ---------------------------------------------------------------------------


class NativeWithStructuredExtractor:
    """Extract text, tables, forms, and images from native structured PDFs."""

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_tables = []
        all_forms = []
        all_images = []
        all_blocks: list[TextBlock] = []

        for page in doc:
            text = page.get_text("text", flags=fitz.TEXT_DEHYPHENATE)
            pages.append(PageResult(page_number=page.number, text=text, method="native"))

            all_tables.extend(_extract_tables_from_page(page))
            all_forms.extend(_extract_form_fields(page))
            all_images.extend(_extract_images_from_page(page))
            all_blocks.extend(_build_text_blocks(page))

        coverage = _text_coverage(pages)
        return ExtractionResult(
            pages=pages,
            method="native_with_structured",
            tables=all_tables,
            forms=all_forms,
            images=all_images,
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="native_with_structured",
                confidence=0.85,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
            ),
        )


# ---------------------------------------------------------------------------
# 3. structured
# ---------------------------------------------------------------------------


class StructuredExtractor:
    """Extract table structures with header relationships and data types."""

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_tables = []
        all_blocks: list[TextBlock] = []

        for page in doc:
            text = page.get_text("text")
            pages.append(PageResult(page_number=page.number, text=text, method="native"))

            tables = _extract_tables_from_page(page)
            all_tables.extend(tables)
            all_blocks.extend(_build_text_blocks(page))

        coverage = _text_coverage(pages)
        return ExtractionResult(
            pages=pages,
            method="structured",
            tables=all_tables,
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="structured",
                confidence=0.8 if all_tables else 0.6,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
            ),
        )


# ---------------------------------------------------------------------------
# 4. scanned_machinewritten
# ---------------------------------------------------------------------------


class ScannedMachinewrittenExtractor:
    """OCR extraction optimised for machine-typed scanned documents."""

    def __init__(self, dpi: int = 200, lang: str = "eng") -> None:
        self.dpi = dpi
        self.lang = lang

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_blocks: list[TextBlock] = []
        all_tables: list[TableData] = []
        confidences: list[float] = []
        combined_steps: list[str] = []

        for page in doc:
            text, conf, steps = _ocr_page(page, self.dpi, self.lang)
            pages.append(PageResult(page_number=page.number, text=text, method="ocr"))
            confidences.append(conf)
            combined_steps.extend(steps)

            # Layout analysis → typed blocks + scanned table structure
            page_blocks, page_tables = _layout_blocks_and_tables(
                page, self.dpi, text, conf,
            )
            all_blocks.extend(page_blocks)
            all_tables.extend(page_tables)

            # Fall back to PyMuPDF table finder if layout didn't find tables
            if not page_tables:
                gray = _page_to_gray(page, dpi=self.dpi)
                from womblex.ingest.heuristics_cv2 import detect_table_grid

                grid = detect_table_grid(gray)
                if grid.has_grid:
                    all_tables.extend(_extract_tables_from_page(page))

        avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
        coverage = _text_coverage(pages)
        unique_steps = sorted(set(combined_steps))

        return ExtractionResult(
            pages=pages,
            method="scanned_machinewritten",
            tables=all_tables,
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="scanned_machinewritten",
                confidence=avg_conf / 100 if avg_conf else 0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
                preprocessing_steps=unique_steps,
            ),
        )


# ---------------------------------------------------------------------------
# 5. scanned_handwritten
# ---------------------------------------------------------------------------


class ScannedHandwrittenExtractor:
    """OCR extraction for handwritten documents with confidence tracking."""

    def __init__(self, dpi: int = 200, lang: str = "eng") -> None:
        self.dpi = dpi
        self.lang = lang

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_blocks: list[TextBlock] = []
        confidences: list[float] = []
        combined_steps: list[str] = []

        for page in doc:
            text, conf, steps = _ocr_page(page, self.dpi, self.lang)
            pages.append(PageResult(page_number=page.number, text=text, method="ocr"))
            confidences.append(conf)
            combined_steps.extend(steps)

            page_blocks, _ = _layout_blocks_and_tables(page, self.dpi, text, conf)
            all_blocks.extend(page_blocks)

        avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
        coverage = _text_coverage(pages)
        unique_steps = sorted(set(combined_steps))

        return ExtractionResult(
            pages=pages,
            method="scanned_handwritten",
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="scanned_handwritten",
                confidence=avg_conf / 100 if avg_conf else 0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
                preprocessing_steps=unique_steps,
            ),
        )


# ---------------------------------------------------------------------------
# 6. scanned_mixed
# ---------------------------------------------------------------------------


class ScannedMixedExtractor:
    """Extract text from documents with both typed and handwritten content."""

    def __init__(self, dpi: int = 200, lang: str = "eng") -> None:
        self.dpi = dpi
        self.lang = lang

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_blocks: list[TextBlock] = []
        all_tables: list[TableData] = []
        confidences: list[float] = []
        combined_steps: list[str] = []
        typed_count = 0
        handwritten_count = 0

        for page in doc:
            gray = _page_to_gray(page, dpi=self.dpi)

            # Classify regions as typed vs handwritten
            from womblex.ingest.heuristics_cv2 import analyze_contour_complexity

            complexity = analyze_contour_complexity(gray)
            is_typed = complexity.regularity > 0.5

            if is_typed:
                typed_count += 1
            else:
                handwritten_count += 1

            text, conf, steps = _ocr_page(page, self.dpi, self.lang)
            pages.append(PageResult(page_number=page.number, text=text, method="ocr"))
            confidences.append(conf)
            combined_steps.extend(steps)

            # Layout analysis for typed blocks and table structure
            page_blocks, page_tables = _layout_blocks_and_tables(
                page, self.dpi, text, conf,
            )
            all_tables.extend(page_tables)

            # Override block_type with typed/handwritten classification
            content_type = "typed" if is_typed else "handwritten"
            for block in page_blocks:
                block = TextBlock(
                    text=block.text,
                    position=block.position,
                    block_type=content_type,
                    confidence=block.confidence,
                )
                all_blocks.append(block)

        avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
        coverage = _text_coverage(pages)
        unique_steps = sorted(set(combined_steps))
        total = typed_count + handwritten_count
        content_mix = {}
        if total > 0:
            content_mix = {
                "typed": typed_count / total,
                "handwritten": handwritten_count / total,
            }

        return ExtractionResult(
            pages=pages,
            method="scanned_mixed",
            tables=all_tables,
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="scanned_mixed",
                confidence=avg_conf / 100 if avg_conf else 0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
                preprocessing_steps=unique_steps,
                content_mix=content_mix,
            ),
        )


# ---------------------------------------------------------------------------
# 7. hybrid
# ---------------------------------------------------------------------------


class HybridExtractor:
    """Extract from documents mixing native text and scanned pages."""

    def __init__(self, dpi: int = 200, lang: str = "eng") -> None:
        self.dpi = dpi
        self.lang = lang

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_tables: list[TableData] = []
        all_forms: list = []
        all_images: list = []
        all_blocks: list[TextBlock] = []
        confidences: list[float] = []
        combined_steps: list[str] = []
        native_count = 0
        ocr_count = 0

        for page in doc:
            native_text = page.get_text("text", flags=fitz.TEXT_DEHYPHENATE).strip()
            is_native = len(native_text) > 100

            if is_native:
                native_count += 1
                pages.append(PageResult(page_number=page.number, text=native_text, method="native"))
                all_tables.extend(_extract_tables_from_page(page))
                all_forms.extend(_extract_form_fields(page))
                all_images.extend(_extract_images_from_page(page))
                all_blocks.extend(_build_text_blocks(page))
                confidences.append(95.0)
            else:
                ocr_count += 1
                text, conf, steps = _ocr_page(page, self.dpi, self.lang)
                pages.append(PageResult(page_number=page.number, text=text, method="ocr"))
                confidences.append(conf)
                combined_steps.extend(steps)

                page_blocks, page_tables = _layout_blocks_and_tables(
                    page, self.dpi, text, conf,
                )
                all_blocks.extend(page_blocks)
                all_tables.extend(page_tables)

        avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
        coverage = _text_coverage(pages)
        unique_steps = sorted(set(combined_steps))
        total = native_count + ocr_count
        content_mix = {}
        if total > 0:
            content_mix = {"native": native_count / total, "scanned": ocr_count / total}

        return ExtractionResult(
            pages=pages,
            method="hybrid",
            tables=all_tables,
            forms=all_forms,
            images=all_images,
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="hybrid",
                confidence=avg_conf / 100 if avg_conf else 0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
                preprocessing_steps=unique_steps,
                content_mix=content_mix,
            ),
        )


# ---------------------------------------------------------------------------
# 8. image
# ---------------------------------------------------------------------------


class ImageExtractor:
    """Extract text and metadata from standalone image files / image PDFs."""

    def __init__(self, dpi: int = 200, lang: str = "eng") -> None:
        self.dpi = dpi
        self.lang = lang

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        pages: list[PageResult] = []
        all_images: list[ImageData] = []
        all_blocks: list[TextBlock] = []
        confidences: list[float] = []
        steps: list[str] = []

        reader = get_paddle_reader(self.lang)

        for page in doc:
            gray = _page_to_gray(page, dpi=self.dpi)

            # Assess blur to determine if OCR is viable
            from womblex.ingest.heuristics_cv2 import calculate_blur_score

            blur = calculate_blur_score(gray)
            if blur is not None and blur < 50:
                steps.append("low_blur_warning")

            # OCR the full page — PaddleOCR returns [(bbox, text, confidence), ...]
            pix = page.get_pixmap(dpi=self.dpi)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            results = reader.readtext(img)
            text = "\n".join(r[1] for r in results if r[1].strip())
            confs = [float(r[2]) for r in results if r[1].strip()]
            avg_conf = (sum(confs) / len(confs)) * 100 if confs else 0.0
            confidences.append(avg_conf)

            pages.append(PageResult(page_number=page.number, text=text, method="ocr"))
            all_images.extend(_extract_images_from_page(page))

            pw, ph = page.rect.width, page.rect.height
            if text.strip():
                all_blocks.append(
                    TextBlock(
                        text=text.strip(),
                        position=_normalise_bbox((0, 0, pw, ph), pw, ph),
                        block_type="paragraph",
                        confidence=avg_conf / 100 if avg_conf else 0.0,
                    )
                )

        avg_conf_doc = sum(confidences) / len(confidences) if confidences else 0.0
        coverage = _text_coverage(pages)
        unique_steps = sorted(set(steps))

        return ExtractionResult(
            pages=pages,
            method="image",
            images=all_images,
            text_blocks=all_blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="image",
                confidence=avg_conf_doc / 100 if avg_conf_doc else 0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=coverage,
                preprocessing_steps=unique_steps,
            ),
        )


# ---------------------------------------------------------------------------
# 9. spreadsheet (CSV / Excel)  — implementation in spreadsheet.py
# ---------------------------------------------------------------------------

# SpreadsheetExtractor is imported at the top of this module and re-exported.
# Its implementation lives in spreadsheet.py to keep this file under 750 lines.


# ---------------------------------------------------------------------------
# 10. docx
# ---------------------------------------------------------------------------


class DocxExtractor:
    """Extract text and tables from Word documents."""

    def extract_path(self, path: Path) -> ExtractionResult:
        """Extract from a DOCX file path (not a fitz.Document)."""
        try:
            from docx import Document
        except ImportError:
            return ExtractionResult(
                pages=[],
                method="docx",
                error="python-docx not installed; cannot extract DOCX.",
                metadata=ExtractionMetadata(
                    extraction_strategy="docx",
                    confidence=0.0,
                    processing_time=0.0,
                    page_count=0,
                    text_coverage=0.0,
                ),
            )

        try:
            doc = Document(str(path))
        except Exception as e:
            return ExtractionResult(
                pages=[],
                method="docx",
                error=f"Failed to read DOCX: {e}",
                metadata=ExtractionMetadata(
                    extraction_strategy="docx",
                    confidence=0.0,
                    processing_time=0.0,
                    page_count=0,
                    text_coverage=0.0,
                ),
            )

        # Extract paragraph text
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # Extract tables
        all_tables = []
        pos = Position(x=0.0, y=0.0, width=1.0, height=1.0)
        for tbl in doc.tables:
            rows_data = []
            for row in tbl.rows:
                rows_data.append([cell.text for cell in row.cells])
            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []
            all_tables.append(TableData(headers=headers, rows=data_rows, position=pos, confidence=0.85))

        # Build text blocks
        blocks = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            block_type = "heading" if para.style and "heading" in para.style.name.lower() else "paragraph"
            blocks.append(TextBlock(text=para.text, position=pos, block_type=block_type, confidence=0.9))

        return ExtractionResult(
            pages=[PageResult(page_number=0, text=text, method="docx")],
            method="docx",
            tables=all_tables,
            text_blocks=blocks,
            metadata=ExtractionMetadata(
                extraction_strategy="docx",
                confidence=0.9,
                processing_time=0.0,
                page_count=1,
                text_coverage=1.0 if text else 0.0,
            ),
        )


# ---------------------------------------------------------------------------
# Non-textual fallback
# ---------------------------------------------------------------------------


class NonTextualExtractor:
    """Placeholder for documents that cannot be extracted -- flags for manual review."""

    def extract(self, doc: fitz.Document) -> ExtractionResult:
        return ExtractionResult(
            pages=[],
            method="non_textual",
            error="Document flagged as non-textual; requires manual review.",
            metadata=ExtractionMetadata(
                extraction_strategy="non_textual",
                confidence=0.0,
                processing_time=0.0,
                page_count=len(doc),
                text_coverage=0.0,
            ),
        )
