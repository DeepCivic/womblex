"""Document type detection for extraction strategy routing.

Profiles a PDF to determine which extraction strategy to use.
Detection is based on text-layer presence, text quality, image presence,
table structure signals, and handwriting indicators.
"""

import logging
import re
import warnings
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

import fitz

from womblex.config import DetectionConfig

# Suppress pymupdf_layout suggestion from find_tables()
warnings.filterwarnings("ignore", message=".*pymupdf_layout.*")

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Document types that drive extraction strategy selection."""

    # PDF types
    SCANNED_HANDWRITTEN = "scanned_handwritten"      # OCR with handwriting config
    SCANNED_MACHINEWRITTEN = "scanned_machinewritten"  # standard OCR
    SCANNED_MIXED = "scanned_mixed"                  # both handwritten and typed
    NATIVE_NARRATIVE = "native_narrative"            # text layer, no structure
    NATIVE_WITH_STRUCTURED = "native_with_structured"  # text layer + tables/images
    STRUCTURED = "structured"                        # pure tabular content
    IMAGE = "image"                                  # photo format (out of scope)
    HYBRID = "hybrid"                                # multiple types in one file
    
    # Non-PDF types
    DOCX = "docx"                                    # Word document (may contain images)
    SPREADSHEET = "spreadsheet"                      # CSV/Excel (may have narrative rows)
    TEXT = "text"                                     # Plain text file (passthrough)
    
    UNKNOWN = "unknown"                              # failed detection


@dataclass
class SheetInfo:
    name: str           # sheet name, or "default" for single-sheet CSV
    sheet_type: str     # "data" | "narrative" | "glossary" | "key_value"
    row_count: int
    col_count: int
    key_column: str | None  # column whose values become part of document_id
    has_sub_headers: bool   # rows that act as section dividers


@dataclass
class DocumentProfile:
    """Result of document type detection."""

    doc_type: DocumentType
    page_count: int
    has_text_layer: bool
    text_coverage: float
    has_images: bool
    has_tables: bool
    has_handwriting_signals: bool
    ocr_confidence: float | None  # None if not sampled
    glyph_regularity: float | None  # 0-1, high = typed (None if not sampled)
    stroke_consistency: float | None  # 0-1, high = typed (None if not sampled)
    confidence: float
    ocr_region_confidences: list[float] | None = None  # per-region scores (0-1) from OCR
    sheet_meta: list[SheetInfo] | None = None  # populated for SPREADSHEET type


# Minimum characters per page to count as having meaningful text.
_MIN_TEXT_LENGTH = 100

# Minimum vector drawing operations to treat a page as "vector-rendered text".
# Pages with many drawings but no text and no images are likely text rendered
# as vector paths (Form XObjects), requiring OCR via pixmap.
_MIN_VECTOR_DRAWINGS = 30

# Pattern for table-like structures: rows with repeated delimiters or whitespace alignment.
_TABLE_PATTERN = re.compile(
    r"(?:"
    r"(?:.*\|.*\|.*\n){2,}"  # pipe-delimited rows
    r"|(?:.*\t.*\t.*\n){2,}"  # tab-delimited rows
    r"|(?:\s{2,}\S+\s{2,}\S+.*\n){3,}"  # whitespace-aligned columns
    r")",
    re.MULTILINE,
)


def _has_table_structure(text: str) -> bool:
    """Detect table-like patterns in extracted text."""
    return bool(_TABLE_PATTERN.search(text))


def _table_signals(
    page: fitz.Page,
    *,
    min_cells: int = 4,
    min_non_empty_cells: int = 300,
    need_manifest: bool = True,
) -> tuple[bool, bool]:
    """Compute ``(has_structural_table, has_manifest_table)`` in one scan.

    PyMuPDF ``find_tables(strategy="text")`` does a full layout analysis and
    costs ~2-3 s on dense pages *regardless of whether it finds a table*.
    ``_has_structural_tables`` and ``_has_manifest_table`` each scan with two
    strategies, so calling both — as the per-page profiler does on every
    table-bearing page — re-runs ``find_tables`` up to 4×. This combined pass
    runs each strategy once and derives both signals, roughly halving per-page
    profiling cost on table-heavy documents (FOI manifests etc.).

    Set ``need_manifest=False`` to skip the per-table ``extract()`` when only
    the structural signal is wanted.

    - Structural: any table (``lines`` or ``text`` strategy) with
      ``row_count * col_count >= min_cells``; the ``text`` strategy needs a
      stricter shape (``rows >= 3, cols >= 2``) to avoid letterhead noise.
    - Manifest: a page dominated by one big table —
      ``extract()`` non-empty cell count ``>= min_non_empty_cells``. Non-empty
      count (not raw cells / empty fraction) is the discriminator that
      separates real manifests (500+ cells/page) from prose-as-table over-fires
      (170-280 cells/page), since the text strategy pads an oversized grid on
      every multi-paragraph page.
    """
    import io
    import sys

    old_stdout = sys.stdout
    sys.stdout = io.StringIO()
    try:
        has_structural = False
        has_manifest = False
        for strategy in ("lines", "text"):
            try:
                tables = page.find_tables(strategy=strategy)
            except Exception:
                continue
            for table in tables.tables:
                # Text-strategy needs stricter shape to avoid letterhead noise.
                if strategy == "text" and (table.row_count < 3 or table.col_count < 2):
                    continue
                if not has_structural and table.row_count * table.col_count >= min_cells:
                    has_structural = True
                if need_manifest and not has_manifest:
                    extracted = table.extract()
                    if extracted:
                        non_empty = sum(
                            1 for row in extracted
                            for cell in row
                            if cell and str(cell).strip()
                        )
                        if non_empty >= min_non_empty_cells:
                            has_manifest = True
                if has_structural and (has_manifest or not need_manifest):
                    return has_structural, has_manifest
        return has_structural, has_manifest
    finally:
        sys.stdout = old_stdout


def _has_structural_tables(page: fitz.Page, min_cells: int = 4) -> bool:
    """Detect tables using PyMuPDF's structural table finder.

    Tries strategy="lines" first (ruled cells), falls back to strategy="text"
    (text-block alignment) so whitespace-aligned columnar layouts like the
    FOI master index get routed through table-aware strategies rather than
    flattened to native_narrative. Thin wrapper over :func:`_table_signals`.
    """
    return _table_signals(page, min_cells=min_cells, need_manifest=False)[0]


def _has_manifest_table(page: fitz.Page, min_non_empty_cells: int = 300) -> bool:
    """Detect manifest-shape tables: a page dominated by one big table.

    Stricter than :func:`_has_structural_tables` — gates the
    ``spreadsheet_print`` qualifier (only manifest-shape pages count toward the
    ≥50%-of-pages rule). Thin wrapper over :func:`_table_signals`; see there for
    the non-empty-cell discriminator rationale.
    """
    return _table_signals(page, min_non_empty_cells=min_non_empty_cells)[1]


def _has_form_structure(page: fitz.Page) -> bool:
    """Detect form field structures on a page.

    Looks for widget annotations (interactive form fields) or
    a high density of short text fragments that suggest labels.
    """
    # Check for interactive form widgets
    widgets = list(page.widgets())
    if len(widgets) >= 2:
        return True

    # Check for label-like text blocks: many short text spans
    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
    short_text_count = 0
    for block in blocks:
        if block.get("type") == 0:  # text block
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if 1 <= len(text) <= 30:
                        short_text_count += 1
    # A page with many short labels is likely a form
    return short_text_count >= 10


# Page-image morphology helpers moved to womblex.ingest.morphology to keep
# detect.py under the 750-line cap. Re-exported below for backward compat.
from womblex.ingest.morphology import (  # noqa: E402,F401
    _analyze_glyph_regularity,
    _analyze_stroke_width_variance,
    _has_handwriting_signals,
    _has_ruled_lines,
    _page_to_grayscale,
    _sample_ocr_confidence,
)




def _classify(
    text_pages: int,
    image_pages: int,
    table_signals: int,
    handwriting_signals: int,
    ocr_confidence: float | None,
    glyph_regularity: float | None,
    stroke_consistency: float | None,
    total_pages: int,
    config: DetectionConfig,
) -> DocumentProfile:
    """Classify a document based on aggregated page-level signals."""
    if total_pages == 0:
        return DocumentProfile(
            doc_type=DocumentType.UNKNOWN,
            page_count=0,
            has_text_layer=False,
            text_coverage=0.0,
            has_images=False,
            has_tables=False,
            has_handwriting_signals=False,
            ocr_confidence=None,
            glyph_regularity=None,
            stroke_consistency=None,
            confidence=0.0,
        )

    text_coverage = text_pages / total_pages
    has_text = text_pages > 0
    has_images = image_pages > 0
    has_tables = table_signals > 0
    has_handwriting = handwriting_signals > 0
    table_ratio = table_signals / total_pages
    handwriting_ratio = handwriting_signals / total_pages if total_pages > 0 else 0.0

    # Native documents (have text layer on most pages)
    if text_coverage >= config.min_text_coverage:
        if table_ratio >= 0.8:
            # Tables on nearly every page — pure tabular content
            doc_type = DocumentType.STRUCTURED
            confidence = 0.85
        elif table_ratio >= config.table_signal_threshold or has_images:
            doc_type = DocumentType.NATIVE_WITH_STRUCTURED
            confidence = 0.85
        else:
            doc_type = DocumentType.NATIVE_NARRATIVE
            confidence = min(0.7 + text_coverage * 0.3, 0.95)
    
    # Hybrid: some pages have text layer, some don't
    elif has_text and has_images and 0.1 < text_coverage < config.min_text_coverage:
        doc_type = DocumentType.HYBRID
        confidence = 0.65
    
    # Scanned documents (no/minimal text layer, need OCR)
    elif has_images:
        # Calculate combined morphology score if available
        morphology_score: float | None = None
        if glyph_regularity is not None and stroke_consistency is not None:
            morphology_score = (glyph_regularity + stroke_consistency) / 2
        elif glyph_regularity is not None:
            morphology_score = glyph_regularity
        elif stroke_consistency is not None:
            morphology_score = stroke_consistency
        
        if handwriting_ratio >= 0.8:
            # Strong handwriting signals (ruled paper detected)
            doc_type = DocumentType.SCANNED_HANDWRITTEN
            confidence = 0.75
        elif has_handwriting and handwriting_ratio < 0.8:
            # Mix of handwritten and typed
            doc_type = DocumentType.SCANNED_MIXED
            confidence = 0.70
        elif morphology_score is not None and morphology_score >= 0.6:
            # High regularity → likely typed/printed
            doc_type = DocumentType.SCANNED_MACHINEWRITTEN
            confidence = min(0.5 + morphology_score * 0.4, 0.85)
        elif morphology_score is not None and morphology_score < 0.35:
            # Low regularity → likely handwritten
            doc_type = DocumentType.SCANNED_HANDWRITTEN
            confidence = 0.6
        elif ocr_confidence is not None and ocr_confidence >= 70:
            # Fall back to OCR confidence if morphology inconclusive
            doc_type = DocumentType.SCANNED_MACHINEWRITTEN
            confidence = min(0.5 + ocr_confidence / 200, 0.85)
        elif ocr_confidence is not None and ocr_confidence < 70:
            # Low OCR confidence → route to UNKNOWN
            doc_type = DocumentType.UNKNOWN
            confidence = 0.4
        else:
            # No morphology/OCR signals available — default to machine-written OCR.
            # If we know there are images, OCR is the right path regardless.
            doc_type = DocumentType.SCANNED_MACHINEWRITTEN
            confidence = 0.5
    
    # Unknown: can't determine
    else:
        doc_type = DocumentType.UNKNOWN
        confidence = 0.3

    return DocumentProfile(
        doc_type=doc_type,
        page_count=total_pages,
        has_text_layer=has_text,
        text_coverage=text_coverage,
        has_images=has_images,
        has_tables=has_tables,
        has_handwriting_signals=has_handwriting,
        ocr_confidence=ocr_confidence,
        glyph_regularity=glyph_regularity,
        stroke_consistency=stroke_consistency,
        confidence=confidence,
    )


def detect_document_type(
    path: Path,
    config: DetectionConfig | None = None,
) -> DocumentProfile:
    """Classify a PDF document for extraction strategy selection.

    Args:
        path: Path to the PDF file.
        config: Detection thresholds. Uses defaults if not provided.

    Returns:
        DocumentProfile with detected type and metadata.
    """
    if config is None:
        config = DetectionConfig()

    doc = fitz.open(str(path))
    try:
        text_pages = 0
        image_pages = 0
        table_signals = 0
        handwriting_signals = 0
        scanned_page_for_analysis: fitz.Page | None = None
        
        total_pages = len(doc)
        max_pages = config.max_sample_pages
        
        # Sample evenly distributed pages up to max_sample_pages
        if total_pages <= max_pages:
            page_indices = list(range(total_pages))
        else:
            step = total_pages / max_pages
            page_indices = [int(i * step) for i in range(max_pages)]

        for idx in page_indices:
            page = doc[idx]
            text = page.get_text().strip()
            images = page.get_images()

            has_meaningful_text = len(text) > _MIN_TEXT_LENGTH
            
            if has_meaningful_text:
                text_pages += 1
                if _has_table_structure(text) or _has_structural_tables(page):
                    table_signals += 1

            # Only count as "image page" (needing OCR) if it has images but lacks text.
            # Native PDFs with logos/graphics still have text layers.
            # Also count pages with heavy vector drawings but no text — these are
            # text rendered as vector paths (Form XObjects) that need OCR via pixmap.
            has_ocr_content = bool(images)
            if not has_meaningful_text and not images:
                drawings = page.get_drawings()
                if len(drawings) >= _MIN_VECTOR_DRAWINGS:
                    has_ocr_content = True

            if has_ocr_content and not has_meaningful_text:
                image_pages += 1
                # Track a scanned page for morphology analysis
                if scanned_page_for_analysis is None:
                    scanned_page_for_analysis = page
                if _has_handwriting_signals(page):
                    handwriting_signals += 1
        
        # Scale counts back to full document estimate
        sample_count = len(page_indices)
        if total_pages > sample_count:
            scale = total_pages / sample_count
            text_pages = int(text_pages * scale)
            image_pages = int(image_pages * scale)
            table_signals = int(table_signals * scale)

        # Sample morphology scores if we have scanned pages
        ocr_confidence: float | None = None
        ocr_region_confidences: list[float] | None = None
        glyph_regularity: float | None = None
        stroke_consistency: float | None = None

        if scanned_page_for_analysis is not None and handwriting_signals == 0:
            # Sample morphology scores (no external binary required)
            glyph_regularity = _analyze_glyph_regularity(scanned_page_for_analysis)
            stroke_consistency = _analyze_stroke_width_variance(scanned_page_for_analysis)

            # Only sample OCR if morphology is inconclusive
            if glyph_regularity is None and stroke_consistency is None:
                ocr_confidence, ocr_region_confidences = _sample_ocr_confidence(
                    scanned_page_for_analysis
                )

        profile = _classify(
            text_pages=text_pages,
            image_pages=image_pages,
            table_signals=table_signals,
            handwriting_signals=handwriting_signals,
            ocr_confidence=ocr_confidence,
            glyph_regularity=glyph_regularity,
            stroke_consistency=stroke_consistency,
            total_pages=total_pages,
            config=config,
        )
        profile.ocr_region_confidences = ocr_region_confidences
        return profile
    finally:
        doc.close()


def _detect_spreadsheet(path: Path) -> DocumentProfile:
    """Inspect a CSV or Excel file and classify each sheet's structure."""
    import pandas as pd
    # Local import: spreadsheet.py → extract.py → detect.py would be circular at module level.
    from womblex.ingest.spreadsheet import (  # noqa: PLC0415
        _HEADER_SCAN_ROWS,
        _classify_sheet,
        read_csv_raw,
        split_preamble,
    )

    # Sample 500 data rows for classification; the read allows headroom
    # for preamble + header rows consumed by split_preamble so the sample
    # size matches the pre-preamble behaviour.
    sample_rows = 500
    read_rows = sample_rows + _HEADER_SCAN_ROWS

    suffix = path.suffix.lower()
    sheet_infos: list[SheetInfo] = []
    try:
        if suffix == ".csv":
            df_raw = read_csv_raw(path, nrows=read_rows)
            _, df = split_preamble(df_raw)
            sheet_infos.append(_classify_sheet("default", df.iloc[:sample_rows]))
        else:
            xl = pd.ExcelFile(str(path))
            for name in xl.sheet_names:
                df_raw = xl.parse(name, dtype=str, keep_default_na=False, nrows=read_rows, header=None)
                _, df = split_preamble(df_raw)
                sheet_infos.append(_classify_sheet(str(name), df.iloc[:sample_rows]))
    except Exception as e:
        logger.warning("spreadsheet detection failed: path=%s error=%s", path, e)

    return DocumentProfile(
        doc_type=DocumentType.SPREADSHEET,
        page_count=len(sheet_infos) or 1,
        has_text_layer=True,
        text_coverage=1.0,
        has_images=False,
        has_tables=True,
        has_handwriting_signals=False,
        ocr_confidence=None,
        glyph_regularity=None,
        stroke_consistency=None,
        confidence=0.9,
        sheet_meta=sheet_infos or None,
    )


def _detect_docx(path: Path) -> DocumentProfile:
    """Detect Word document characteristics.

    Note: DOCX files may contain embedded images that need OCR.
    This basic detection doesn't analyze image content.
    """
    has_images = False
    has_tables = False
    text_length = 0
    try:
        from docx import Document
        doc = Document(str(path))
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_images = True
                break
        has_tables = len(doc.tables) > 0
        for para in doc.paragraphs:
            text_length += len(para.text)
    except ImportError:
        # python-docx not installed, return basic profile
        pass
    except Exception:
        # Failed to parse, return unknown
        return DocumentProfile(
            doc_type=DocumentType.UNKNOWN,
            page_count=0,
            has_text_layer=False,
            text_coverage=0.0,
            has_images=False,
            has_tables=False,
            has_handwriting_signals=False,
            ocr_confidence=None,
            glyph_regularity=None,
            stroke_consistency=None,
            confidence=0.3,
        )
    
    return DocumentProfile(
        doc_type=DocumentType.DOCX,
        page_count=1,  # DOCX doesn't expose page count easily
        has_text_layer=True,
        text_coverage=1.0 if text_length > 0 else 0.0,
        has_images=has_images,
        has_tables=has_tables,
        has_handwriting_signals=False,
        ocr_confidence=None,
        glyph_regularity=None,
        stroke_consistency=None,
        confidence=0.85,
    )


def detect_file_type(
    path: Path,
    config: DetectionConfig | None = None,
) -> DocumentProfile:
    """Classify any supported file for extraction strategy selection.

    Handles PDFs, Word documents, spreadsheets, and plain text files.

    Args:
        path: Path to the file.
        config: Detection thresholds. Uses defaults if not provided.

    Returns:
        DocumentProfile with detected type and metadata.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return detect_document_type(path, config)
    elif suffix == ".docx":
        return _detect_docx(path)
    elif suffix in (".csv", ".xlsx", ".xls"):
        return _detect_spreadsheet(path)
    elif suffix == ".txt":
        return DocumentProfile(
            doc_type=DocumentType.TEXT,
            page_count=1,
            has_text_layer=True,
            text_coverage=1.0,
            has_images=False,
            has_tables=False,
            has_handwriting_signals=False,
            ocr_confidence=None,
            glyph_regularity=None,
            stroke_consistency=None,
            confidence=1.0,
        )
    else:
        return DocumentProfile(
            doc_type=DocumentType.UNKNOWN,
            page_count=0,
            has_text_layer=False,
            text_coverage=0.0,
            has_images=False,
            has_tables=False,
            has_handwriting_signals=False,
            ocr_confidence=None,
            glyph_regularity=None,
            stroke_consistency=None,
            confidence=0.0,
        )
